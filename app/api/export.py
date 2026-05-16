import io
import json
from pathlib import Path
from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import FileResponse, Response
from app.database.db import get_db
from app.config import TRANSCRIPTS_DIR

router = APIRouter()


def _fmt_ts(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def _fmt_duration_hms(seconds: float) -> str:
    return _fmt_ts(seconds)


def _build_transcript_data(recording_id: int) -> dict:
    with get_db() as conn:
        rec = conn.execute("SELECT * FROM recordings WHERE id = ?", (recording_id,)).fetchone()
        if not rec:
            raise HTTPException(status_code=404, detail="Grabación no encontrada.")

        segs = conn.execute(
            "SELECT s.start_time, s.end_time, s.text, sp.display_name, s.raw_speaker_label "
            "FROM segments s "
            "LEFT JOIN speakers sp ON s.speaker_id = sp.id "
            "WHERE s.recording_id = ? ORDER BY s.start_time",
            (recording_id,),
        ).fetchall()

        participants_rows = conn.execute(
            "SELECT DISTINCT sp.display_name FROM segments s "
            "JOIN speakers sp ON s.speaker_id = sp.id "
            "WHERE s.recording_id = ?",
            (recording_id,),
        ).fetchall()

    participants = [r["display_name"] for r in participants_rows if r["display_name"]]
    segments = [
        {
            "start": seg["start_time"],
            "end": seg["end_time"],
            "text": seg["text"],
            "speaker": seg["display_name"] or seg["raw_speaker_label"] or "Desconocido",
        }
        for seg in segs
    ]

    speaker_times: dict[str, float] = {}
    for seg in segments:
        name = seg["speaker"]
        speaker_times[name] = speaker_times.get(name, 0.0) + (seg["end"] - seg["start"])

    total_duration = rec["duration_seconds"] or sum(s["end"] - s["start"] for s in segments)

    return {
        "filename": rec["filename"],
        "created_at": rec["created_at"],
        "completed_at": rec["completed_at"],
        "duration": total_duration,
        "language": rec["language_detected"] or "Desconocido",
        "participants": participants,
        "segments": segments,
        "speaker_times": speaker_times,
        "segment_count": len(segments),
    }


def _to_txt(data: dict) -> str:
    lines = [
        "TRANSCRIPCIÓN DE REUNIÓN",
        "========================",
        f"Título: {data['filename']}",
        f"Fecha: {data['completed_at'] or data['created_at'] or 'Desconocida'}",
        f"Duración: {_fmt_duration_hms(data['duration'])}",
        f"Participantes: {', '.join(data['participants']) or 'No identificados'}",
        f"Idioma: {data['language']}",
        "Procesado con: VozMeet",
        "========================",
        "",
    ]
    for seg in data["segments"]:
        lines.append(f"[{_fmt_ts(seg['start'])}] {seg['speaker']}")
        lines.append(seg["text"])
        lines.append("")

    total = data["duration"]
    lines += [
        "========================",
        "FIN DE TRANSCRIPCIÓN",
        f"Total de intervenciones: {data['segment_count']}",
        "Tiempo hablado por persona:",
    ]
    for name, t in sorted(data["speaker_times"].items(), key=lambda x: -x[1]):
        pct = round(t / total * 100, 1) if total else 0
        m, s = divmod(int(t), 60)
        lines.append(f"  {name}: {m} min {s} seg ({pct}%)")
    lines.append("========================")

    return "\n".join(lines)


def _to_md(data: dict) -> str:
    lines = [
        f"# Transcripción: {data['filename']}",
        "",
        f"**Fecha:** {data['completed_at'] or data['created_at'] or 'Desconocida'}  ",
        f"**Duración:** {_fmt_duration_hms(data['duration'])}  ",
        f"**Participantes:** {', '.join(data['participants']) or 'No identificados'}  ",
        f"**Idioma:** {data['language']}  ",
        f"**Procesado con:** VozMeet",
        "",
        "---",
        "",
    ]
    for seg in data["segments"]:
        lines.append(f"**[{_fmt_ts(seg['start'])}] {seg['speaker']}**")
        lines.append(f"> {seg['text']}")
        lines.append("")

    lines += ["---", "", "## Resumen de participación", ""]
    total = data["duration"]
    for name, t in sorted(data["speaker_times"].items(), key=lambda x: -x[1]):
        pct = round(t / total * 100, 1) if total else 0
        m, s = divmod(int(t), 60)
        lines.append(f"- **{name}**: {m} min {s} seg ({pct}%)")

    return "\n".join(lines)


def _to_docx(data: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="python-docx no instalado. Ejecuta: pip install python-docx"
        )

    doc = Document()

    title = doc.add_heading(f"Transcripción: {data['filename']}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta.add_run(f"Fecha: ").bold = True
    meta.add_run(f"{data['completed_at'] or data['created_at'] or 'Desconocida'}\n")
    meta.add_run(f"Duración: ").bold = True
    meta.add_run(f"{_fmt_duration_hms(data['duration'])}\n")
    meta.add_run(f"Participantes: ").bold = True
    meta.add_run(f"{', '.join(data['participants']) or 'No identificados'}\n")
    meta.add_run(f"Idioma: ").bold = True
    meta.add_run(f"{data['language']}\n")
    meta.add_run(f"Procesado con: ").bold = True
    meta.add_run("VozMeet")

    doc.add_paragraph()
    doc.add_heading("Transcripción", level=1)

    spk_colors = [
        RGBColor(0, 122, 255),
        RGBColor(255, 59, 48),
        RGBColor(52, 199, 89),
        RGBColor(255, 149, 0),
        RGBColor(175, 82, 222),
        RGBColor(0, 199, 190),
    ]
    spk_color_map = {}
    color_idx = 0

    for seg in data["segments"]:
        spk = seg["speaker"]
        if spk not in spk_color_map:
            spk_color_map[spk] = spk_colors[color_idx % len(spk_colors)]
            color_idx += 1

        p = doc.add_paragraph()
        ts_run = p.add_run(f"[{_fmt_ts(seg['start'])}] ")
        ts_run.font.size = Pt(9)
        ts_run.font.color.rgb = RGBColor(0x8E, 0x8E, 0x93)

        spk_run = p.add_run(f"{spk}  ")
        spk_run.bold = True
        spk_run.font.color.rgb = spk_color_map[spk]

        p.add_run(seg["text"] or "")
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()
    doc.add_heading("Participación", level=1)
    total = data["duration"]
    for name, t in sorted(data["speaker_times"].items(), key=lambda x: -x[1]):
        pct = round(t / total * 100, 1) if total else 0
        m, s = divmod(int(t), 60)
        doc.add_paragraph(f"{name}: {m} min {s} seg ({pct}%)", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/export/{recording_id}")
def export_transcript(
    recording_id: int,
    format: str = Query("txt", regex="^(txt|md|json|docx)$"),
):
    data = _build_transcript_data(recording_id)
    stem = Path(data["filename"]).stem[:40]

    if format == "json":
        output_bytes = json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")
        filename = f"{stem}_vozmeet.json"
        content_type = "application/json"
    elif format == "md":
        output_bytes = _to_md(data).encode("utf-8")
        filename = f"{stem}_vozmeet.md"
        content_type = "text/markdown; charset=utf-8"
    elif format == "docx":
        output_bytes = _to_docx(data)
        filename = f"{stem}_vozmeet.docx"
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        output_bytes = _to_txt(data).encode("utf-8")
        filename = f"{stem}_vozmeet.txt"
        content_type = "text/plain; charset=utf-8"

    out_path = TRANSCRIPTS_DIR / filename
    out_path.write_bytes(output_bytes)

    return FileResponse(
        path=str(out_path),
        media_type=content_type,
        filename=filename,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
